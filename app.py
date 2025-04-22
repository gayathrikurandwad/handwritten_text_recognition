import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import pandas as pd
from tqdm import tqdm
from mltu.configs import BaseModelConfigs
import shutil
import os
import cv2
import typing
import numpy as np
import docx
from spellchecker import SpellChecker
from docx.shared import Pt
import checkout
from mltu.inferenceModel import OnnxInferenceModel
from mltu.utils.text_utils import ctc_decoder, get_cer
from flask_classful import FlaskView
from flask import Flask,request,render_template,send_file
from db import db_init,db
from werkzeug.utils import secure_filename
from models import Img,Files
import sqlite3
from io import BytesIO
from PIL import Image
from docx2pdf import convert

spell = SpellChecker()
app=Flask(__name__)
# Define the function to upload and save the image
# def upload_image():
# 	file_path = filedialog.askopenfilename()
# 	if file_path:
# 		image = Image.open(file_path)
# 		image.thumbnail((300, 300)) # Resize image if necessary
# 		photo = ImageTk.PhotoImage(image)
# 		image_label.config(image=photo)
# 		image_label.image = photo # Keep a reference to avoid garbage collection
# 		save_image(file_path)
# 		messagebox.showinfo("Success", "Image uploaded successfully!")
#
# def save_image(file_path):
# 	save_dir = "saved_images"
# 	if not os.path.exists(save_dir):
# 		os.makedirs(save_dir)
# 	filename = os.path.basename(file_path)
# 	shutil.copy(file_path, os.path.join(save_dir, "1.png"))
# 	print("Image saved to:", os.path.join(save_dir, "1.png"))
#
# # Create the main window
# root = tk.Tk()
# root.title("Image Uploader")
#
# # Set window size
# root.geometry("500x500")
#
# # Create and pack widgets
# upload_button = tk.Button(root, text="Upload Image", command=upload_image)
# upload_button.pack(pady=10)
#
# image_label = tk.Label(root)
# image_label.pack()
#
# # Run the application
# root.mainloop()




class ImageToWordModel(OnnxInferenceModel):
    def __init__(self, char_list: typing.Union[str, list], *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.char_list = char_list

    def predict(self, image: np.ndarray):
        image = cv2.resize(image, self.input_shapes[0][1:3][::-1])

        image_pred = np.expand_dims(image, axis=0).astype(np.float32)

        preds = self.model.run(self.output_names, {self.input_names[0]: image_pred})[0]

        text = ctc_decoder(preds, self.char_list)[0]

        return text

if __name__ == "__main__":
    app.run()



app.config['SQLALCHEMY_DATABASE_URI']='sqlite:///data.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS']=False
db_init(app)

@app.route('/')
@app.route('/index')
def index():
    return render_template('index.html') # index.html


@app.route('/upload')
def pic():
    return render_template('upload.html')

@app.route('/uploading',methods=['POST'])
def upload():
    pic=request.files['pic']
    if not pic:
        return "No pic uploaded",400
    
    filename=secure_filename(pic.filename)
    mimetype=pic.mimetype
    img=Img(img=pic.read(),mimetype=mimetype,name=filename)
    db.session.add(img)
    db.session.commit()
    image=img.img
    images=Image.open(BytesIO(image))
    images.save('static/images/3.png')
    return render_template('detect.html'),200
    
@app.route('/detect',methods=['POST'])
def main_img_word():
    configs = BaseModelConfigs.load("Models/word_recog/202410042200/configs.yaml")

    model = ImageToWordModel(model_path=configs.model_path, char_list=configs.vocab)

    df = pd.read_csv("Models/word_recog/202410042200/val.csv").values.tolist()
    doc = docx.Document()
    para = doc.add_paragraph()
    dir_path='segmented'
    count=0
    for f in os.listdir(dir_path):
        fp=os.path.join(dir_path,f)
        if os.path.isfile(fp):
            os.remove(fp)
    checkout.segment_para()
    for path in os.listdir(dir_path):
        if os.path.isfile(os.path.join(dir_path,path)):
            count+=1
    for i in range(count):
        image = cv2.imread(f"segmented/segment{i}.png")
        word= model.predict(image)
        prediction_text=spell.correction(word)
        para.add_run(f"{prediction_text} ")
        doc.save('static/documents/detected.docx')
        print(f"Prediction: {prediction_text}")
        doc=docx.Document("static/documents/detected.docx")
        full_text=[]
        for para in doc.paragraphs:
            full_text.append(para.text)
        line1=" ".join(full_text)
    return render_template('download.html',text=line1)

def save_file_to_db(file_path):
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    filename = os.path.basename(file_path)
    with open(file_path, "rb") as file:
        file_data = file.read()
    document = Files(name=filename, file=file_data)
    db.session.add(document)
    db.session.commit()

@app.route('/download',methods=['POST'])
def download():
    file_path="static/documents/detected.docx"
    save_file_to_db(file_path)
    up=db.session.query(Files).order_by(Files.id.desc()).first()
    return send_file(BytesIO(up.file),download_name=up.name,as_attachment=True)

@app.route('/pdf',methods=['POST'])
def pdf():
    convert("static/documents/detected.docx")
    return send_file("static/documents/detected.pdf",as_attachment=True)

    